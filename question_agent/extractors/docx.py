"""Word .docx text extractor using python-docx."""

import io
from typing import Any

import docx

from question_agent.extractors.exceptions import ExtractionError


def extract_docx(content: bytes) -> dict[str, Any]:
    """Extract text from .docx bytes.  Returns dict with text, paragraph_count, table_count."""
    try:
        document = docx.Document(io.BytesIO(content))
    except Exception as e:
        raise ExtractionError(
            "File is not a valid .docx document. "
            "Legacy .doc files are not supported — please convert to .docx first.",
            format="docx",
        ) from e

    paragraphs: list[str] = []
    for para in document.paragraphs:
        style = para.style.name if para.style and para.style.name != "Normal" else None
        prefix = f"[{style}] " if style else ""
        t = para.text.strip()
        if t:
            paragraphs.append(f"{prefix}{t}")

    table_texts: list[str] = []
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                table_texts.append(" | ".join(cells))

    parts = []
    if paragraphs:
        parts.append("\n".join(paragraphs))
    if table_texts:
        parts.append("\n\n[Tables]\n" + "\n".join(table_texts))

    return {
        "text": "\n\n".join(parts) if parts else "",
        "paragraph_count": len(paragraphs),
        "table_count": len(document.tables),
    }


def extract_docx_structured(content: bytes) -> dict[str, Any]:
    """Extract structured paragraphs from .docx with style_name per paragraph.

    Returns paragraphs as list of dicts with text and style_name (no [Style] prefix).
    Normal style is omitted (style_name = None).
    """
    try:
        document = docx.Document(io.BytesIO(content))
    except Exception as e:
        raise ExtractionError(
            "File is not a valid .docx document. "
            "Legacy .doc files are not supported — please convert to .docx first.",
            format="docx",
        ) from e

    paragraphs: list[dict[str, Any]] = []
    for para in document.paragraphs:
        t = para.text.strip()
        if not t:
            continue
        style = para.style.name if para.style and para.style.name != "Normal" else None
        paragraphs.append({"text": t, "style_name": style})

    char_count = sum(len(p["text"]) for p in paragraphs)

    return {
        "paragraphs": paragraphs,
        "char_count": char_count,
        "paragraph_count": len(paragraphs),
        "table_count": len(document.tables),
    }
