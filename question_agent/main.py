"""FastAPI application entry point."""

import io
from collections.abc import AsyncGenerator
from contextlib import asynccontextmanager

import docx
import pdfplumber
from fastapi import FastAPI, File, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse
from pydantic import BaseModel
from starlette import status

from question_agent import __version__
from question_agent.config import settings

MAX_UPLOAD_BYTES = 50 * 1024 * 1024  # 50 MB

SUPPORTED_FORMATS = {
    ".txt": "text",
    ".pdf": "pdf",
    ".docx": "docx",
}

MIME_TO_FORMAT = {
    "text/plain": "text",
    "application/pdf": "pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    "application/octet-stream": None,  # ambiguous, fall back to extension
}


class ExtractResponse(BaseModel):
    format: str
    text: str
    page_count: int | None = None
    paragraph_count: int | None = None
    table_count: int | None = None
    warnings: list[str] | None = None


def _detect_format(filename: str, content_type: str | None) -> str | None:
    """Detect file format from extension and MIME type."""
    # Try MIME first
    if content_type and content_type in MIME_TO_FORMAT:
        fmt = MIME_TO_FORMAT[content_type]
        if fmt is not None:
            return fmt

    # Fall back to extension
    import os

    ext = os.path.splitext(filename)[1].lower()
    return SUPPORTED_FORMATS.get(ext)


@asynccontextmanager
async def lifespan(app: FastAPI) -> AsyncGenerator[None, None]:
    """Application lifespan — startup and shutdown events."""
    # Startup
    yield
    # Shutdown


app = FastAPI(
    title="AI 智能出题智能体",
    description=(
        "基于 GLM-5 的智能出题引擎 — 支持知识点抽取、题干生成、干扰项构建、答案解析与质量评估"
    ),
    version=__version__,
    lifespan=lifespan,
    docs_url="/docs" if settings.debug else None,
    redoc_url="/redoc" if settings.debug else None,
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


@app.api_route("/health", methods=["GET", "HEAD"])
async def health_check() -> JSONResponse:
    """Health check endpoint."""
    return JSONResponse(
        content={
            "status": "healthy",
            "version": __version__,
        }
    )


@app.post("/extract/text")
async def extract_text(file: UploadFile = File(...)) -> JSONResponse:
    """Extract plain text from an uploaded .txt file."""
    if file.content_type and not file.content_type.startswith("text/"):
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail="Only text/plain files are accepted.",
        )

    content = await file.read()

    if len(content) > MAX_UPLOAD_BYTES:
        raise HTTPException(
            status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
            detail="File size exceeds 50 MB limit.",
        )

    text = content.decode("utf-8", errors="replace")
    return JSONResponse(content={"text": text})


@app.post("/extract/pdf")
async def extract_pdf(file: UploadFile = File(...)) -> JSONResponse:
    """Extract text from an uploaded PDF file, page by page."""
    if file.content_type and file.content_type != "application/pdf":
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail="Only application/pdf files are accepted.",
        )

    content = await file.read()

    if len(content) > MAX_UPLOAD_BYTES:
        raise HTTPException(
            status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
            detail="File size exceeds 50 MB limit.",
        )

    try:
        pdf = pdfplumber.open(io.BytesIO(content))
    except Exception:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail="File is not a valid PDF or is corrupted.",
        )

    pages = []
    warnings: list[str] = []
    total_pages = len(pdf.pages)
    for i, page in enumerate(pdf.pages):
        text = page.extract_text(layout=True)
        if text:
            pages.append(text)
        else:
            warnings.append(f"Page {i + 1}: no extractable text layer (likely scanned image)")

    pdf.close()

    full_text = "\n--- PAGE BREAK ---\n".join(pages)

    return JSONResponse(
        content={
            "text": full_text,
            "page_count": total_pages,
            "pages_with_text": len(pages),
            "warnings": warnings or None,
        }
    )


DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


@app.post("/extract/docx")
async def extract_docx(file: UploadFile = File(...)) -> JSONResponse:
    """Extract text from an uploaded .docx file, including paragraphs and tables."""
    if file.content_type and file.content_type not in (DOCX_MIME, "application/octet-stream"):
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail=(
                f"Expected .docx format ({DOCX_MIME}), got: {file.content_type}. "
                "Legacy .doc files are not supported."
            ),
        )

    content = await file.read()

    if len(content) > MAX_UPLOAD_BYTES:
        raise HTTPException(
            status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
            detail="File size exceeds 50 MB limit.",
        )

    try:
        document = docx.Document(io.BytesIO(content))
    except Exception:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail=(
                "File is not a valid .docx document. "
                "Legacy .doc files are not supported — please convert to .docx first."
            ),
        )

    paragraphs: list[str] = []
    for para in document.paragraphs:
        style = para.style.name if para.style and para.style.name != "Normal" else None
        prefix = f"[{style}] " if style else ""
        text = para.text.strip()
        if text:
            paragraphs.append(f"{prefix}{text}")

    table_texts: list[str] = []
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                table_texts.append(" | ".join(cells))

    parts = []
    if paragraphs:
        parts.append("\n".join(paragraphs))
    if table_texts:
        parts.append("\n\n[Tables]\n" + "\n".join(table_texts))

    return JSONResponse(
        content={
            "text": "\n\n".join(parts) if parts else "",
            "paragraph_count": len(paragraphs),
            "table_count": len(document.tables),
        }
    )


@app.post("/extract")
async def extract_unified(file: UploadFile = File(...)) -> JSONResponse:
    """Unified extraction endpoint — auto-detects format and extracts text."""
    fmt = _detect_format(file.filename or "", file.content_type)

    if fmt is None:
        supported = ", ".join(SUPPORTED_FORMATS.keys())
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail=f"Unsupported file format. Supported extensions: {supported}",
        )

    content = await file.read()

    if len(content) > MAX_UPLOAD_BYTES:
        raise HTTPException(
            status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
            detail="File size exceeds 50 MB limit.",
        )

    if fmt == "text":
        text = content.decode("utf-8", errors="replace")
        return JSONResponse(content=ExtractResponse(format="text", text=text).model_dump())

    elif fmt == "pdf":
        try:
            pdf = pdfplumber.open(io.BytesIO(content))
        except Exception:
            raise HTTPException(
                status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                detail="File is not a valid PDF or is corrupted.",
            )
        pages = []
        warnings: list[str] = []
        total_pages = len(pdf.pages)
        for i, page in enumerate(pdf.pages):
            text = page.extract_text(layout=True)
            if text:
                pages.append(text)
            else:
                warnings.append(f"Page {i + 1}: no extractable text layer (likely scanned image)")
        pdf.close()
        return JSONResponse(
            content=ExtractResponse(
                format="pdf",
                text="\n--- PAGE BREAK ---\n".join(pages),
                page_count=total_pages,
                warnings=warnings or None,
            ).model_dump()
        )

    elif fmt == "docx":
        try:
            document = docx.Document(io.BytesIO(content))
        except Exception:
            raise HTTPException(
                status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                detail=(
                    "File is not a valid .docx document. "
                    "Legacy .doc files are not supported — please convert to .docx first."
                ),
            )
        paragraphs: list[str] = []
        for para in document.paragraphs:
            style = para.style.name if para.style and para.style.name != "Normal" else None
            prefix = f"[{style}] " if style else ""
            t = para.text.strip()
            if t:
                paragraphs.append(f"{prefix}{t}")
        table_texts: list[str] = []
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    table_texts.append(" | ".join(cells))
        parts = []
        if paragraphs:
            parts.append("\n".join(paragraphs))
        if table_texts:
            parts.append("\n\n[Tables]\n" + "\n".join(table_texts))
        return JSONResponse(
            content=ExtractResponse(
                format="docx",
                text="\n\n".join(parts) if parts else "",
                paragraph_count=len(paragraphs),
                table_count=len(document.tables),
            ).model_dump()
        )

    # Should be unreachable
    raise HTTPException(status_code=status.HTTP_422_UNPROCESSABLE_ENTITY, detail="Unknown format")


def main() -> None:
    """Entry point for `question-agent` CLI script."""
    import uvicorn

    uvicorn.run(
        "question_agent.main:app",
        host=settings.host,
        port=settings.port,
        reload=settings.debug,
    )


if __name__ == "__main__":
    main()
