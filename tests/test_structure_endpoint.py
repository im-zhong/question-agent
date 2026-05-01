"""Tests for the /structure endpoint."""

import io

import docx
import pytest
from fpdf import FPDF
from httpx import AsyncClient


@pytest.mark.asyncio
async def test_structure_text_chinese_numbering(client: AsyncClient):
    """POST /structure on Chinese-numbered text returns chapter tree."""
    text = "第一章 引言\n这是引言内容。\n1.1 背景\n背景详细信息。\n第二章 方法\n方法描述。"
    resp = await client.post(
        "/structure",
        files={"file": ("test.txt", text.encode(), "text/plain")},
    )
    assert resp.status_code == 200
    data = resp.json()
    assert data["format"] == "text"
    assert data["chapters"] is not None
    assert len(data["chapters"]) >= 2
    assert data["detection_stats"]["total"] >= 2
    assert data["detection_stats"]["method"] in ("hybrid", "rule_only")


@pytest.mark.asyncio
async def test_structure_docx(client: AsyncClient):
    """POST /structure on DOCX returns chapter tree."""
    doc = docx.Document()
    doc.add_heading("Chapter 1", level=1)
    doc.add_paragraph("Body text.")
    doc.add_heading("Section 1.1", level=2)
    doc.add_paragraph("More text.")
    doc.add_heading("Chapter 2", level=1)
    doc.add_paragraph("Final text.")

    buf = io.BytesIO()
    doc.save(buf)
    docx_bytes = buf.getvalue()

    resp = await client.post(
        "/structure",
        files={
            "file": (
                "test.docx",
                docx_bytes,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert resp.status_code == 200
    data = resp.json()
    assert data["format"] == "docx"
    assert data["chapters"] is not None
    assert len(data["chapters"]) >= 2
    # Verify tree structure
    for ch in data["chapters"]:
        assert "id" in ch
        assert "level" in ch
        assert "title" in ch
        assert "children" in ch


@pytest.mark.asyncio
async def test_structure_pdf(client: AsyncClient):
    """POST /structure on PDF returns chapter tree."""
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", size=16)
    pdf.cell(text="Chapter 1")
    pdf.ln()
    pdf.set_font("Helvetica", size=10)
    pdf.cell(text="Body text here.")
    pdf_bytes = bytes(pdf.output())

    resp = await client.post(
        "/structure",
        files={"file": ("test.pdf", pdf_bytes, "application/pdf")},
    )
    assert resp.status_code == 200
    data = resp.json()
    assert data["format"] == "pdf"
    assert "chapters" in data
    assert "detection_stats" in data


@pytest.mark.asyncio
async def test_structure_no_headings(client: AsyncClient):
    """POST /structure on text with no headings returns chapters=null."""
    text = "This is just plain text.\nNo headings here.\nJust content."
    resp = await client.post(
        "/structure",
        files={"file": ("test.txt", text.encode(), "text/plain")},
    )
    assert resp.status_code == 200
    data = resp.json()
    assert data["detection_stats"]["total"] == 0


@pytest.mark.asyncio
async def test_structure_unsupported_format(client: AsyncClient):
    """POST /structure with unsupported format returns 422."""
    resp = await client.post(
        "/structure",
        files={"file": ("test.jpg", b"fake", "image/jpeg")},
    )
    assert resp.status_code == 422


@pytest.mark.asyncio
async def test_health_still_works(client: AsyncClient):
    """/health endpoint is not affected by /structure."""
    resp = await client.get("/health")
    assert resp.status_code == 200
    assert resp.json()["status"] == "healthy"
