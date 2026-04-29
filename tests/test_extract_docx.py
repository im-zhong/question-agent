"""Tests for F-4: Word 文本提取."""

import io

import docx
import httpx
import pytest

BASE = "http://localhost:8000"


def _make_docx(paragraphs: list[str], heading: str | None = None) -> bytes:
    """Build a minimal .docx with given paragraphs."""
    doc = docx.Document()
    if heading:
        doc.add_heading(heading, level=1)
    for p in paragraphs:
        doc.add_paragraph(p)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _make_docx_with_table() -> bytes:
    """Build a .docx with a table."""
    doc = docx.Document()
    doc.add_paragraph("Intro paragraph.")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "A"
    table.cell(0, 1).text = "B"
    table.cell(1, 0).text = "C"
    table.cell(1, 1).text = "D"
    doc.save(buf := io.BytesIO())
    return buf.getvalue()


DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


@pytest.mark.asyncio
async def test_extract_docx_returns_paragraphs():
    docx_bytes = _make_docx(["Para one", "Para two"])
    async with httpx.AsyncClient() as client:
        resp = await client.post(
            f"{BASE}/extract/docx",
            files={"file": ("test.docx", docx_bytes, DOCX_MIME)},
        )
        assert resp.status_code == 200
        data = resp.json()
        assert "Para one" in data["text"]
        assert "Para two" in data["text"]
        assert data["paragraph_count"] == 2


@pytest.mark.asyncio
async def test_extract_docx_heading_style_annotated():
    docx_bytes = _make_docx(["Body text"], heading="Chapter Title")
    async with httpx.AsyncClient() as client:
        resp = await client.post(
            f"{BASE}/extract/docx",
            files={"file": ("test.docx", docx_bytes, DOCX_MIME)},
        )
        assert resp.status_code == 200
        data = resp.json()
        assert "[Heading 1]" in data["text"]
        assert "Chapter Title" in data["text"]


@pytest.mark.asyncio
async def test_extract_docx_extracts_table():
    docx_bytes = _make_docx_with_table()
    async with httpx.AsyncClient() as client:
        resp = await client.post(
            f"{BASE}/extract/docx",
            files={"file": ("test.docx", docx_bytes, DOCX_MIME)},
        )
        assert resp.status_code == 200
        data = resp.json()
        assert "[Tables]" in data["text"]
        assert "A | B" in data["text"]
        assert "C | D" in data["text"]
        assert data["table_count"] == 1


@pytest.mark.asyncio
async def test_extract_docx_invalid_file_returns_422():
    async with httpx.AsyncClient() as client:
        resp = await client.post(
            f"{BASE}/extract/docx",
            files={"file": ("bad.docx", b"not a zip file", DOCX_MIME)},
        )
        assert resp.status_code == 422
        assert "Legacy .doc" in resp.json()["detail"]


@pytest.mark.asyncio
async def test_extract_docx_wrong_mime_returns_422():
    docx_bytes = _make_docx(["test"])
    async with httpx.AsyncClient() as client:
        resp = await client.post(
            f"{BASE}/extract/docx",
            files={"file": ("test.docx", docx_bytes, "text/plain")},
        )
        assert resp.status_code == 422


@pytest.mark.asyncio
async def test_extract_docx_missing_file_returns_422():
    async with httpx.AsyncClient() as client:
        resp = await client.post(f"{BASE}/extract/docx")
        assert resp.status_code == 422


@pytest.mark.asyncio
async def test_extract_docx_octet_stream_accepted():
    """Generic application/octet-stream MIME should be accepted."""
    docx_bytes = _make_docx(["Content"])
    async with httpx.AsyncClient() as client:
        resp = await client.post(
            f"{BASE}/extract/docx",
            files={"file": ("test.docx", docx_bytes, "application/octet-stream")},
        )
        assert resp.status_code == 200
