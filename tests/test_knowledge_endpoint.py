"""Tests for POST /knowledge endpoint."""

import pytest
from httpx import ASGITransport, AsyncClient

from question_agent.main import app

BASE = "http://test"


@pytest.fixture
async def client() -> AsyncClient:
    transport = ASGITransport(app=app)
    async with AsyncClient(transport=transport, base_url=BASE) as c:
        yield c


def _make_txt(content: str = "加速度是指速度变化量与时间的比值。公式：v = v₀ + at") -> bytes:
    return content.encode("utf-8")


def _make_docx() -> bytes:
    from docx import Document

    doc = Document()
    doc.add_heading("第一章 力学基础", level=1)
    doc.add_paragraph("牛顿第二定律 F=ma，描述了力与加速度的关系。")
    doc.add_paragraph("加速度的定义：速度变化量与时间的比值。")
    from io import BytesIO

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _make_pdf() -> bytes:
    from fpdf import FPDF

    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", size=12)
    pdf.cell(text="Newton's Second Law F=ma")
    pdf.ln()
    pdf.cell(text="Acceleration is the rate of change of velocity.")
    from io import BytesIO

    buf = BytesIO()
    pdf.output(buf)
    return buf.getvalue()


class TestKnowledgeEndpoint:
    @pytest.mark.asyncio
    async def test_txt_upload(self, client: AsyncClient) -> None:
        resp = await client.post(
            f"{BASE}/knowledge",
            files={"file": ("test.txt", _make_txt(), "text/plain")},
        )
        assert resp.status_code == 200
        data = resp.json()
        assert data["format"] == "text"
        assert "knowledge_points" in data
        assert "extraction_stats" in data
        assert "total" in data["extraction_stats"]
        assert "method" in data["extraction_stats"]

    @pytest.mark.asyncio
    async def test_docx_upload(self, client: AsyncClient) -> None:
        resp = await client.post(
            f"{BASE}/knowledge",
            files={
                "file": (
                    "test.docx",
                    _make_docx(),
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )
        assert resp.status_code == 200
        data = resp.json()
        assert data["format"] == "docx"
        assert "knowledge_points" in data

    @pytest.mark.asyncio
    async def test_pdf_upload(self, client: AsyncClient) -> None:
        resp = await client.post(
            f"{BASE}/knowledge",
            files={"file": ("test.pdf", _make_pdf(), "application/pdf")},
        )
        assert resp.status_code == 200
        data = resp.json()
        assert data["format"] == "pdf"

    @pytest.mark.asyncio
    async def test_no_knowledge_points_document(self, client: AsyncClient) -> None:
        resp = await client.post(
            f"{BASE}/knowledge",
            files={"file": ("plain.txt", b"Hello world, nothing special here.", "text/plain")},
        )
        assert resp.status_code == 200
        data = resp.json()
        # May have 0 knowledge points (no markers, no LLM in test)
        assert "knowledge_points" in data
        assert isinstance(data["knowledge_points"], list)

    @pytest.mark.asyncio
    async def test_knowledge_point_response_structure(self, client: AsyncClient) -> None:
        resp = await client.post(
            f"{BASE}/knowledge",
            files={"file": ("test.txt", _make_txt(), "text/plain")},
        )
        assert resp.status_code == 200
        data = resp.json()
        if data["knowledge_points"]:
            kp = data["knowledge_points"][0]
            assert "id" in kp
            assert "name" in kp
            assert "description" in kp
            assert "tags" in kp
            assert "confidence" in kp
            assert "method" in kp
            assert "source_line_start" in kp
            assert "source_line_end" in kp
            assert "chapter_id" in kp

    @pytest.mark.asyncio
    async def test_unsupported_format_returns_422(self, client: AsyncClient) -> None:
        resp = await client.post(
            f"{BASE}/knowledge",
            files={"file": ("test.xyz", b"data", "application/octet-stream")},
        )
        assert resp.status_code == 422


class TestExistingEndpointsUnchanged:
    @pytest.mark.asyncio
    async def test_health_still_works(self, client: AsyncClient) -> None:
        resp = await client.get(f"{BASE}/health")
        assert resp.status_code == 200
        assert resp.json()["status"] == "healthy"

    @pytest.mark.asyncio
    async def test_extract_still_works(self, client: AsyncClient) -> None:
        resp = await client.post(
            f"{BASE}/extract",
            files={"file": ("test.txt", b"Hello", "text/plain")},
        )
        assert resp.status_code == 200

    @pytest.mark.asyncio
    async def test_structure_still_works(self, client: AsyncClient) -> None:
        resp = await client.post(
            f"{BASE}/structure",
            files={"file": ("test.txt", b"Hello world", "text/plain")},
        )
        assert resp.status_code == 200
