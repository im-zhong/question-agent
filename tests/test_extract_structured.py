"""Tests for structured extraction mode (?mode=structured)."""

import io

import docx
import pytest
from fpdf import FPDF
from httpx import AsyncClient


@pytest.mark.asyncio
async def test_structured_pdf_paragraph_metadata(client: AsyncClient):
    """POST /extract?mode=structured on PDF returns per-paragraph font metadata."""
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", size=16)
    pdf.cell(text="Chapter Title")
    pdf.ln()
    pdf.set_font("Helvetica", style="B", size=12)
    pdf.cell(text="Bold Section")
    pdf.ln()
    pdf.set_font("Helvetica", size=10)
    pdf.cell(text="Normal text here.")
    pdf_bytes = bytes(pdf.output())

    resp = await client.post(
        "/extract?mode=structured",
        files={"file": ("test.pdf", pdf_bytes, "application/pdf")},
    )

    assert resp.status_code == 200
    data = resp.json()
    assert data["format"] == "pdf"
    assert len(data["paragraphs"]) >= 3

    # All paragraphs should have font metadata fields
    for p in data["paragraphs"]:
        assert "text" in p
        assert "page_number" in p
        assert "font_size" in p
        assert "is_bold" in p
        assert "x0" in p

    # At least one bold paragraph
    bold_paragraphs = [p for p in data["paragraphs"] if p["is_bold"]]
    assert len(bold_paragraphs) >= 1


@pytest.mark.asyncio
async def test_structured_docx_style_name(client: AsyncClient):
    """POST /extract?mode=structured on DOCX returns style_name, no [Heading] prefix."""
    doc = docx.Document()
    doc.add_heading("Chapter 1", level=1)
    doc.add_paragraph("Some body text.")
    doc.add_heading("Section 1.1", level=2)
    doc.add_paragraph("More content.")

    buf = io.BytesIO()
    doc.save(buf)
    docx_bytes = buf.getvalue()

    resp = await client.post(
        "/extract?mode=structured",
        files={
            "file": (
                "test.docx",
                docx_bytes,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )

    assert resp.status_code == 200
    data = resp.json()
    assert data["format"] == "docx"

    headings = [p for p in data["paragraphs"] if p.get("style_name")]
    assert len(headings) >= 2

    # Verify no [Style] prefix in text
    for p in data["paragraphs"]:
        assert not p["text"].startswith("[")
        assert "style_name" in p

    # Verify specific headings
    h1 = [p for p in data["paragraphs"] if p.get("style_name") == "Heading 1"]
    assert len(h1) == 1
    assert h1[0]["text"] == "Chapter 1"


@pytest.mark.asyncio
async def test_structured_text_null_font_fields(client: AsyncClient):
    """POST /extract?mode=structured on plain text returns null font fields."""
    text_content = "Line one\nLine two\nLine three"

    resp = await client.post(
        "/extract?mode=structured",
        files={"file": ("test.txt", text_content.encode(), "text/plain")},
    )

    assert resp.status_code == 200
    data = resp.json()
    assert data["format"] == "text"
    assert len(data["paragraphs"]) == 3

    for p in data["paragraphs"]:
        assert "text" in p
        assert p.get("font_size") is None
        assert p.get("is_bold") is None
        assert p.get("x0") is None
        assert p.get("page_number") is None
        assert p.get("style_name") is None


@pytest.mark.asyncio
async def test_structured_text_empty_file(client: AsyncClient):
    """POST /extract?mode=structured on empty text file returns empty paragraphs."""
    resp = await client.post(
        "/extract?mode=structured",
        files={"file": ("empty.txt", b"", "text/plain")},
    )

    assert resp.status_code == 200
    data = resp.json()
    assert data["format"] == "text"
    assert data["paragraphs"] == []
    assert data["char_count"] == 0


@pytest.mark.asyncio
async def test_mode_text_unchanged_pdf(client: AsyncClient):
    """POST /extract?mode=text on PDF returns flat text (same as before)."""
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", size=12)
    pdf.cell(text="Hello World")
    pdf_bytes = bytes(pdf.output())

    resp = await client.post(
        "/extract?mode=text",
        files={"file": ("test.pdf", pdf_bytes, "application/pdf")},
    )

    assert resp.status_code == 200
    data = resp.json()
    assert data["format"] == "pdf"
    assert "text" in data
    assert "paragraphs" not in data
    assert "Hello World" in data["text"]


@pytest.mark.asyncio
async def test_default_mode_unchanged_text(client: AsyncClient):
    """POST /extract (no mode param) returns flat text (backward compat)."""
    resp = await client.post(
        "/extract",
        files={"file": ("test.txt", b"hello world", "text/plain")},
    )

    assert resp.status_code == 200
    data = resp.json()
    assert data["format"] == "text"
    assert "text" in data
    assert "paragraphs" not in data
    assert data["text"] == "hello world"


@pytest.mark.asyncio
async def test_structured_invalid_mode_rejected(client: AsyncClient):
    """POST /extract?mode=invalid returns 422."""
    resp = await client.post(
        "/extract?mode=invalid",
        files={"file": ("test.txt", b"test", "text/plain")},
    )

    assert resp.status_code == 422


@pytest.mark.asyncio
async def test_structured_text_single_line(client: AsyncClient):
    """POST /extract?mode=structured on single line text."""
    resp = await client.post(
        "/extract?mode=structured",
        files={"file": ("test.txt", b"single line", "text/plain")},
    )

    assert resp.status_code == 200
    data = resp.json()
    assert len(data["paragraphs"]) == 1
    assert data["paragraphs"][0]["text"] == "single line"
